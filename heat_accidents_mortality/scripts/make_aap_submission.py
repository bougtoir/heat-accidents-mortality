#!/usr/bin/env python3
"""
Convert the Lancet Planetary Health manuscript into an Accident Analysis &
Prevention (AAP) submission package.

The script re-uses make_manuscript.py (which reads all numbers from the
processed result CSVs) to regenerate fresh inline and legends-only Word files,
then post-processes them into AAP format:
  - unstructured abstract + keywords
  - removal of the Research-in-context panel
  - numbered main sections (1. Introduction ... 5. Conclusion)
  - superscript {1,2} citations converted to bracketed [1,2] citations

It also writes an AAP cover letter, an editable highlights DOCX, and bundles the
final submission package (manuscript, highlights, cover letter, tables,
figures, STROBE, and separate figure files).
"""
import os
import re
import shutil
import sys
import tempfile
import zipfile

from docx import Document
from docx.shared import Pt

# make_manuscript.py lives in the same scripts/ directory
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
MAN = os.path.join(ROOT, "output", "manuscript")
FIG = os.path.join(ROOT, "output", "figures")

sys.path.insert(0, HERE)
import make_manuscript as mm  # noqa: E402


# ---------------------------------------------------------------------------
# AAP formatting constants
# ---------------------------------------------------------------------------
AAP_KEYWORDS = "heat; traffic crash; mortality; climate change; distributed-lag model; road safety"

HIGHLIGHTS = [
    "Hotter-than-normal days are linked to same-day US traffic-crash mortality.",
    "The excess is comparable in magnitude to all recorded direct-heat deaths.",
    "Open-air road users show the largest heat-attributable excess risk.",
    "Uniform +1 to +3 °C warming would add hundreds of deaths per year.",
]

AAP_COVER = """[PLACEHOLDER date]

The Editors, Accident Analysis & Prevention

Dear Editors,

We submit for your consideration our manuscript, "Hotter-than-normal days and traffic-crash mortality: a distributed-lag analysis of the United States and Japan and the question of under-recognised heat illness", as a Research Article.

Using only public data, this ecological time-series study shows that days hotter than the local seasonal norm carry an acute excess of US traffic-crash deaths. The excess survives adjustment for national driving activity, is concentrated in heat-exposed open-air road users (motorcyclists, pedestrians and cyclists), and is comparable in magnitude to all officially recorded direct-heat deaths (ICD-10 X30). Scenario projections suggest this burden would grow under uniform warming. These findings are directly relevant to the journal's scope of transportation-accident injury and the environmental and human factors that influence crash occurrence and severity.

The manuscript is original, is not under consideration elsewhere, and all authors approve submission. All data are public and the complete analysis pipeline is openly available and fully reproducible (make all for data and figures, then make aap for the manuscript and submission package), with no hard-coded results. We declare [PLACEHOLDER competing interests]. We confirm [PLACEHOLDER preprint status].

We look forward to your assessment.

Yours sincerely,

[PLACEHOLDER corresponding author, on behalf of all authors]
"""


def _punct_move_citations(doc):
    """Convert Word-font-superscript {1,2} markers to bracketed [1,2].

    Punctuation that immediately precedes the superscript run is moved to after
    the bracket, consistent with Vancouver style using square brackets.
    """
    punct = set(".,;:!?")
    for p in doc.paragraphs:
        runs = p.runs
        for i, r in enumerate(runs):
            if not r.font.superscript or not r.text:
                continue
            nums = r.text
            prev = runs[i - 1].text if i > 0 else ""
            trailing = ""
            if prev and prev[-1] in punct:
                trailing = prev[-1]
                runs[i - 1].text = prev[:-1]
                leading = " "
            elif prev and prev[-1].isspace():
                leading = ""
            elif not prev or not prev.strip():
                leading = ""
            else:
                leading = " "
            r.text = f"{leading}[{nums}]{trailing}"
            r.font.superscript = None


def _heading_text(p):
    return (p.text or "").strip()


def _set_heading_text(p, text):
    # Preserve Heading style; clear runs and set single run with new text.
    p.text = text


def _remove_paragraph(p):
    p._element.getparent().remove(p._element)


def _repurpose_as_keywords(p, keywords):
    """Replace paragraph contents with a bold 'Keywords: ' label + keyword text."""
    p.text = ""
    r1 = p.add_run("Keywords: ")
    r1.bold = True
    p.add_run(keywords)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(6)


def _make_abstract(doc, abstract_heading, body_ps):
    """Merge labelled Background/Methods/Findings/Interpretation paragraphs."""
    label_map = {
        "Background": 0,
        "Methods": 1,
        "Findings": 2,
        "Interpretation": 3,
    }
    ordered = [None, None, None, None]
    funding_p = None
    for p in body_ps:
        txt = p.text or ""
        for lab in label_map:
            if txt.startswith(lab):
                ordered[label_map[lab]] = re.sub(rf"^{lab}\s*", "", txt)
                break
        else:
            if txt.startswith("Funding"):
                funding_p = p
    parts = [t for t in ordered if t is not None]
    abstract_text = " ".join(parts)

    # The first body paragraph (Background) becomes the single abstract paragraph.
    if body_ps:
        body_ps[0].text = abstract_text
        body_ps[0].paragraph_format.line_spacing = 2.0
        body_ps[0].paragraph_format.space_after = Pt(6)

    # Remove the remaining labelled body paragraphs (except Funding, which we
    # repurpose as the Keywords paragraph).
    for p in body_ps[1:]:
        if p is funding_p:
            continue
        _remove_paragraph(p)

    if funding_p:
        _repurpose_as_keywords(funding_p, AAP_KEYWORDS)

    _set_heading_text(abstract_heading, "Abstract")


def _remove_research_in_context(doc, heading, body_ps):
    _remove_paragraph(heading)
    for p in body_ps:
        _remove_paragraph(p)


def _renumber_headings(doc):
    renames = {
        "Introduction": "1. Introduction",
        "Methods": "2. Methods",
        "Results": "3. Results",
        "Discussion": "4. Discussion",
        "Conclusion": "5. Conclusion",
    }
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            txt = _heading_text(p)
            if txt in renames:
                _set_heading_text(p, renames[txt])


def _postprocess(src_path, dst_path):
    doc = Document(src_path)

    # 1. Citations in text.
    _punct_move_citations(doc)

    # 2. Locate section headings and abstract/research-in-context bodies.
    abstract_heading = None
    abstract_body = []
    research_heading = None
    research_body = []

    collect = None
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            txt = _heading_text(p)
            if txt == "Summary":
                abstract_heading = p
                collect = "abstract"
                continue
            if txt == "Research in context":
                research_heading = p
                collect = "research"
                continue
            # Any other heading terminates collection.
            collect = None
            continue
        if collect == "abstract":
            abstract_body.append(p)
        elif collect == "research":
            research_body.append(p)

    # 3. Restructure abstract / keywords.
    if abstract_heading and abstract_body:
        _make_abstract(doc, abstract_heading, abstract_body)

    # 4. Drop Research-in-context panel.
    if research_heading:
        _remove_research_in_context(doc, research_heading, research_body)

    # 5. Renumber main IMRAD sections.
    _renumber_headings(doc)

    doc.save(dst_path)
    print("wrote", dst_path)


def _build_highlights(path):
    doc = Document()
    mm.setup(doc)
    t = doc.add_paragraph()
    t.alignment = mm.WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Highlights")
    r.bold = True
    r.font.size = Pt(14)
    for bullet in HIGHLIGHTS:
        p = doc.add_paragraph(bullet)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        # Use a bullet style if available; otherwise the leading disc is fine.
        try:
            p.style = doc.styles["List Bullet"]
        except Exception:
            pass
    doc.save(path)
    print("wrote", path)


def _build_cover_letter(path):
    doc = Document()
    mm.setup(doc)
    for block in AAP_COVER.strip().split("\n\n"):
        text = " ".join(block.splitlines())
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
    doc.save(path)
    print("wrote", path)


def _build_figures_temp():
    """Copy the existing submission figure files into a temporary staging folder."""
    dest_dir = tempfile.mkdtemp(prefix="aap_figures_", dir=MAN)
    src = os.path.join(MAN, "submission_figures")
    if not os.path.isdir(src):
        raise FileNotFoundError(f"{src} not found; run 'make manuscript' first.")
    shutil.copytree(src, os.path.join(dest_dir, "aap_submission_figures"))
    return dest_dir


def _build_submission_zip(fig_stage):
    zip_base = os.path.join(MAN, "aap_submission_package")
    stage = zip_base + "_stage"
    if os.path.exists(stage):
        shutil.rmtree(stage)
    os.makedirs(stage)
    files = [
        "heat_crash_mortality_aap.docx",
        "heat_crash_mortality_aap_legends.docx",
        "aap_highlights.docx",
        "aap_cover_letter.docx",
        "tables.docx",
        "figures.pptx",
        "strobe_checklist.docx",
    ]
    for name in files:
        src = os.path.join(MAN, name)
        if os.path.exists(src):
            shutil.copyfile(src, os.path.join(stage, name))
    if fig_stage and os.path.isdir(fig_stage):
        for entry in os.listdir(fig_stage):
            entry_src = os.path.join(fig_stage, entry)
            entry_dst = os.path.join(stage, entry)
            if os.path.isdir(entry_src):
                shutil.copytree(entry_src, entry_dst)
            else:
                shutil.copyfile(entry_src, entry_dst)
    if os.path.exists(zip_base + ".zip"):
        os.remove(zip_base + ".zip")
    shutil.make_archive(zip_base, "zip", stage)
    shutil.rmtree(stage)
    print("wrote", zip_base + ".zip")


def main():
    os.makedirs(MAN, exist_ok=True)

    # Generate fresh base manuscripts in a temporary directory so the existing
    # tracked Lancet-style files are not overwritten with new timestamps.
    tmpdir = tempfile.mkdtemp(prefix="aap_build_", dir=MAN)
    original_man = mm.MAN
    mm.MAN = tmpdir
    try:
        mm.build_manuscript("tmp_inline.docx", embed=True)
        mm.build_manuscript("tmp_legends.docx", embed=False)
    finally:
        mm.MAN = original_man

    inline_src = os.path.join(tmpdir, "tmp_inline.docx")
    legends_src = os.path.join(tmpdir, "tmp_legends.docx")

    aap_inline = os.path.join(MAN, "heat_crash_mortality_aap.docx")
    aap_legends = os.path.join(MAN, "heat_crash_mortality_aap_legends.docx")

    _postprocess(inline_src, aap_inline)
    _postprocess(legends_src, aap_legends)

    # Clean up temporary base files.
    shutil.rmtree(tmpdir)

    _build_highlights(os.path.join(MAN, "aap_highlights.docx"))
    _build_cover_letter(os.path.join(MAN, "aap_cover_letter.docx"))
    fig_stage = _build_figures_temp()
    try:
        _build_submission_zip(fig_stage)
    finally:
        shutil.rmtree(fig_stage)

    # Character-count sanity check for highlights.
    for i, h in enumerate(HIGHLIGHTS, 1):
        if len(h) > 85:
            print(f"WARNING: highlight {i} exceeds 85 characters ({len(h)})")


if __name__ == "__main__":
    main()
