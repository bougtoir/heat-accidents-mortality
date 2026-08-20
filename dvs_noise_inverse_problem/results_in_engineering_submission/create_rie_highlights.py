#!/usr/bin/env python3
"""
Generate Highlights document for Results in Engineering (Elsevier) submission.

Elsevier requirements:
- Separate source file (Microsoft Word, not PDF)
- 3 to 5 bullet points
- Each highlight max 85 characters including spaces
- File name: "Highlights"
"""

import json
from docx import Document
from docx.shared import Pt
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
OUT_DIR = SCRIPT_DIR


def load_metrics():
    results_dir = SCRIPT_DIR.parent / 'results'
    with open(results_dir / 'evaluation_summary.json') as f:
        summary = json.load(f)
    with open(results_dir / 'demo_summary.json') as f:
        demo = json.load(f)
    m = summary['methods']
    return {
        'fano_auc': m['fano_filter']['auc_mean'],
        'fano_spr': m['fano_filter']['spr_mean'],
        'demo_removal': demo['removal_pct'],
    }


def build_highlights():
    metrics = load_metrics()
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    p = doc.add_paragraph()
    run = p.add_run('Highlights')
    run.font.size = Pt(14)
    run.bold = True

    highlights = [
        'Closed-form optimal noise-removal level for event-based sensors is derived.',
        'Five-parameter circuit-level noise model (A5) inverts DVS pixel statistics.',
        f"Fano-factor-based filter achieves ROC-AUC {metrics['fano_auc']:.3f} while preserving {metrics['fano_spr']*100:.1f}% signal.",
        'Six-tier calibration hierarchy includes satellite-trail verification tier.',
        f"Open pipeline demonstrates {metrics['demo_removal']}% noise removal and trajectory recovery.",
    ]

    for h in highlights:
        assert len(h) <= 85, f"Highlight too long ({len(h)} chars): {h}"
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    out_path = OUT_DIR / 'highlights_rie.docx'
    doc.save(str(out_path))
    print(f"Highlights saved: {out_path}")

    for i, h in enumerate(highlights, 1):
        print(f"  [{i}] {len(h)} chars: {h}")

    return out_path


if __name__ == '__main__':
    build_highlights()
