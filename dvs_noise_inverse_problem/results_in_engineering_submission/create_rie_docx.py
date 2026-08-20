#!/usr/bin/env python3
"""
Generate integrated manuscript for Results in Engineering (Elsevier, ISSN 2590-1230).

Title: "Noise removal for dynamic vision sensors:
        a physics-informed stochastic resonance framework"

Integrates:
  - Paper a: DVS noise inverse problem (JATIS submission)
  - Paper b: Covariate-adjusted stochastic resonance (PRE submission)

Results in Engineering formatting:
- Single-anonymised review -> blinded manuscript (author info on separate title page)
- Separate title page (create_rie_title_page.py)
- Numbered references in square brackets [1], [1,2], [1-4]
- Times New Roman 12pt, double-spaced
- Figures inline with captions below
- Highlights (3-5 bullet points) on first page
- No word limit; engineering relevance and reproducibility are emphasised
- Equations as Word equation objects (OMML)

Numerical values are loaded from results/evaluation_summary.json and
results/demo_summary.json so the manuscript stays reproducible.
"""

import json
import re
import sys
from lxml import etree
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
OUT_DIR = SCRIPT_DIR
# In-repo figure directories (public mirror root is dvs_noise_inverse_problem)
FIG_DIR_DVS = SCRIPT_DIR.parent          # noise_inverse_demo.py / systematic_evaluation.py outputs
FIG_DIR_SR = SCRIPT_DIR / 'sr_figures'   # generate_sr_figures.py outputs

# Global metrics cache; populated by load_results(). Keys may be referenced
# with {key} placeholders in manuscript strings so numerical values are not
# hard-coded.
_METRICS = {}


# =========================================================
# OMML equation helpers (reused from JATIS script)
# =========================================================

def _mr(parent, text, italic=True, bold=False):
    r = etree.SubElement(parent, qn('m:r'))
    if not italic or bold:
        rPr = etree.SubElement(r, qn('m:rPr'))
        if not italic:
            sty = etree.SubElement(rPr, qn('m:sty'))
            sty.set(qn('m:val'), 'p')
        if bold:
            sty = etree.SubElement(rPr, qn('m:sty'))
            sty.set(qn('m:val'), 'b')
    t = etree.SubElement(r, qn('m:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r


def _sub(parent, base, sub):
    el = etree.SubElement(parent, qn('m:sSub'))
    e = etree.SubElement(el, qn('m:e'))
    _mr(e, base)
    s = etree.SubElement(el, qn('m:sub'))
    _mr(s, sub)
    return el


def _sup(parent, base, sup):
    el = etree.SubElement(parent, qn('m:sSup'))
    e = etree.SubElement(el, qn('m:e'))
    _mr(e, base)
    s = etree.SubElement(el, qn('m:sup'))
    _mr(s, sup)
    return el


def _sup_builder(parent, base_builder, sup_text):
    """Superscript where the base is built by a callable (e.g. delimited fraction)."""
    el = etree.SubElement(parent, qn('m:sSup'))
    e = etree.SubElement(el, qn('m:e'))
    base_builder(e)
    s = etree.SubElement(el, qn('m:sup'))
    _mr(s, sup_text)
    return el


def _frac(parent, num_builder, den_builder):
    f = etree.SubElement(parent, qn('m:f'))
    num = etree.SubElement(f, qn('m:num'))
    num_builder(num)
    den = etree.SubElement(f, qn('m:den'))
    den_builder(den)
    return f


def _delim(parent, content_builder, left='(', right=')'):
    d = etree.SubElement(parent, qn('m:d'))
    dPr = etree.SubElement(d, qn('m:dPr'))
    begChr = etree.SubElement(dPr, qn('m:begChr'))
    begChr.set(qn('m:val'), left)
    endChr = etree.SubElement(dPr, qn('m:endChr'))
    endChr.set(qn('m:val'), right)
    e = etree.SubElement(d, qn('m:e'))
    content_builder(e)
    return d


def _hat(parent, text):
    acc = etree.SubElement(parent, qn('m:acc'))
    accPr = etree.SubElement(acc, qn('m:accPr'))
    chrEl = etree.SubElement(accPr, qn('m:chr'))
    chrEl.set(qn('m:val'), '\u0302')
    e = etree.SubElement(acc, qn('m:e'))
    _mr(e, text)
    return acc


def _sqrt(parent, content_builder):
    rad = etree.SubElement(parent, qn('m:rad'))
    radPr = etree.SubElement(rad, qn('m:radPr'))
    degHide = etree.SubElement(radPr, qn('m:degHide'))
    degHide.set(qn('m:val'), '1')
    deg = etree.SubElement(rad, qn('m:deg'))
    e = etree.SubElement(rad, qn('m:e'))
    content_builder(e)
    return rad


def add_display_equation(doc, builder_func, eq_num=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    omathpara = etree.SubElement(p._element, qn('m:oMathPara'))
    omath = etree.SubElement(omathpara, qn('m:oMath'))
    builder_func(omath)
    if eq_num:
        run = p.add_run(f'    ({eq_num})')
        run.font.size = Pt(12)
    return p


# =========================================================
# Equation definitions for integrated manuscript
# =========================================================

def eq_sr_snr(omath):
    """Eq: SNR_out(sigma) proportional to (A/sigma^2)^2 exp(-2 theta^2 / sigma^2)"""
    _sub(omath, 'SNR', 'out')
    def _arg(e):
        _mr(e, '\u03c3')
    _delim(omath, _arg)
    _mr(omath, ' \u221d ')
    def _frac_delim(e):
        def _n(n):
            _mr(n, 'A')
        def _d(d):
            _sup(d, '\u03c3', '2')
        _frac(e, _n, _d)
    _sup_builder(omath, _frac_delim, '2')
    _mr(omath, ' ')
    _mr(omath, 'exp', italic=False)
    def _exp_arg(e):
        _mr(e, '\u22122')
        _frac(e,
               lambda n: _sup(n, '\u03b8', '2'),
               lambda d: _sup(d, '\u03c3', '2'))
    _delim(omath, _exp_arg)


def eq_sigma_eff(omath):
    """Eq: sigma_eff = sigma * sqrt(1 - rho^2)"""
    _sub(omath, '\u03c3', 'eff')
    _mr(omath, ' = \u03c3')
    def _sq(e):
        _mr(e, '1 \u2212 ')
        _sup(e, '\u03c1', '2')
    _sqrt(omath, _sq)


def eq_rho_star(omath):
    """Eq: rho* = sqrt(1 - theta^2 / sigma^2), for sigma > theta"""
    _sup(omath, '\u03c1', '*')
    _mr(omath, ' = ')
    def _sq(e):
        _mr(e, '1 \u2212 ')
        _frac(e,
               lambda n: _sup(n, '\u03b8', '2'),
               lambda d: _sup(d, '\u03c3', '2'))
    _sqrt(omath, _sq)
    _mr(omath, ', \u03c3 > \u03b8', italic=False)


def eq_snr_gain(omath):
    """Eq: SNR gain = (sigma/theta)^4 exp(2(sigma^2 - theta^2)/sigma^2)"""
    def _n(n):
        _sub(n, 'SNR', 'out')
        def _arg(e):
            _mr(e, '\u03c3, ')
            _sup(e, '\u03c1', '*')
        _delim(n, _arg)
    def _d(d):
        _sub(d, 'SNR', 'out')
        def _arg(e):
            _mr(e, '\u03c3, 0')
        _delim(d, _arg)
    _frac(omath, _n, _d)
    _mr(omath, ' = ')
    def _frac_delim(e):
        def _n2(n):
            _mr(n, '\u03c3')
        def _d2(d):
            _mr(d, '\u03b8')
        _frac(e, _n2, _d2)
    _sup_builder(omath, _frac_delim, '4')
    _mr(omath, ' ')
    _mr(omath, 'exp', italic=False)
    def _exp_arg(e):
        _mr(e, '2')
        def _n3(n):
            _sup(n, '\u03c3', '2')
            _mr(n, ' \u2212 ')
            _sup(n, '\u03b8', '2')
        def _d3(d):
            _sup(d, '\u03c3', '2')
        _frac(e, _n3, _d3)
    _delim(omath, _exp_arg)


def eq_a5_model(omath):
    """Eq: lambda_noise(T, I_bg) = I_dark,ref * exp(alpha * DeltaT) * (1 + beta * I_bg)"""
    _sub(omath, '\u03bb', 'noise')
    def _args(e):
        _mr(e, 'T, ')
        _sub(e, 'I', 'bg')
    _delim(omath, _args)
    _mr(omath, ' = ')
    _sub(omath, 'I', 'dark,ref')
    _mr(omath, ' \u22c5 exp')
    def _exp_arg(e):
        _mr(e, '\u03b1 \u22c5 \u0394T')
    _delim(omath, _exp_arg)
    _mr(omath, ' \u22c5 ')
    def _bg(e):
        _mr(e, '1 + \u03b2 \u22c5 ')
        _sub(e, 'I', 'bg')
    _delim(omath, _bg)


def eq_accuracy(omath):
    """Eq: alpha = 1 - ||e_hat - e_true|| / ||e_true||"""
    _mr(omath, '\u03b1 = 1 \u2212 ')
    def _num(n):
        _mr(n, '\u2016')
        _hat(n, 'e')
        _sub(n, '', 'noise')
        _mr(n, ' \u2212 ')
        _sub(n, 'e', 'noise,true')
        _mr(n, '\u2016')
    def _den(d):
        _mr(d, '\u2016')
        _sub(d, 'e', 'noise,true')
        _mr(d, '\u2016')
    _frac(omath, _num, _den)


def eq_residual_noise(omath):
    """Eq: sigma_residual = (1 - alpha) * sigma_original"""
    _sub(omath, '\u03c3', 'residual')
    _mr(omath, ' = ')
    def _factor(e):
        _mr(e, '1 \u2212 \u03b1')
    _delim(omath, _factor)
    _mr(omath, ' \u22c5 ')
    _sub(omath, '\u03c3', 'original')


def eq_snr_improvement(omath):
    """Eq: SNR_after / SNR_before = 1 / (1 - alpha)"""
    def _num(n):
        _sub(n, 'SNR', 'after')
    def _den(d):
        _sub(d, 'SNR', 'before')
    _frac(omath, _num, _den)
    _mr(omath, ' = ')
    def _num2(n):
        _mr(n, '1', italic=False)
    def _den2(d):
        _mr(d, '1 \u2212 \u03b1')
    _frac(omath, _num2, _den2)


def eq_map_estimation(omath):
    """Eq: theta_hat = argmax p(e_cal|theta) * p(theta|theta_prior)"""
    _hat(omath, '\u03b8')
    _mr(omath, ' = ')
    _sub(omath, 'argmax', '\u03b8')
    _mr(omath, ' p')
    def _likelihood(e):
        _sub(e, 'e', 'cal')
        _mr(e, ' | \u03b8')
    _delim(omath, _likelihood)
    _mr(omath, ' \u22c5 p')
    def _prior(e):
        _mr(e, '\u03b8 | ')
        _sub(e, '\u03b8', 'prior')
    _delim(omath, _prior)


def eq_nn_output(omath):
    """Eq: lambda_hat_noise = lambda_physics * (1 + Delta_lambda_aux) + Delta_lambda_corr"""
    _hat(omath, '\u03bb')
    _sub(omath, '', 'noise')
    _mr(omath, ' = ')
    _sub(omath, '\u03bb', 'physics')
    _mr(omath, ' \u22c5 ')
    def _mod(e):
        _mr(e, '1 + \u0394')
        _sub(e, '\u03bb', 'aux')
    _delim(omath, _mod)
    _mr(omath, ' + \u0394')
    _sub(omath, '\u03bb', 'corr')


def eq_p_noise(omath):
    """Eq: P_noise(e_i) = lambda_hat_noise / (lambda_hat_noise + lambda_hat_signal)"""
    _sub(omath, 'P', 'noise')
    def _ei(e):
        _sub(e, 'e', 'i')
    _delim(omath, _ei)
    _mr(omath, ' = ')
    def _num(n):
        _hat(n, '\u03bb')
        _sub(n, '', 'noise')
        def _coords(e):
            _sub(e, 'x', 'i')
            _mr(e, ', ')
            _sub(e, 'y', 'i')
            _mr(e, ', ')
            _sub(e, 't', 'i')
        _delim(n, _coords)
    def _den(d):
        _hat(d, '\u03bb')
        _sub(d, '', 'noise')
        def _c1(e):
            _mr(e, '\u2026')
        _delim(d, _c1)
        _mr(d, ' + ')
        _hat(d, '\u03bb')
        _sub(d, '', 'signal')
        def _c2(e):
            _mr(e, '\u2026')
        _delim(d, _c2)
    _frac(omath, _num, _den)


def eq_fano(omath):
    """Eq: F = Var(N_k) / Mean(N_k)"""
    _mr(omath, 'F', italic=True)
    _mr(omath, ' = ', italic=False)
    def _num(n):
        _mr(n, 'Var', italic=False)
        def _nk(e):
            _sub(e, 'N', 'k')
        _delim(n, _nk)
    def _den(d):
        _mr(d, 'Mean', italic=False)
        def _nk(e):
            _sub(e, 'N', 'k')
        _delim(d, _nk)
    _frac(omath, _num, _den)


def eq_detection_limit(omath):
    """Eq: Delta_m approx 2.5 log10(1/(1-alpha))"""
    _mr(omath, '\u0394m \u2248 2.5 ')
    _sub(omath, 'log', '10')
    def _arg(e):
        def _n(n):
            _mr(n, '1', italic=False)
        def _d(d):
            _mr(d, '1 \u2212 \u03b1')
        _frac(e, _n, _d)
    _delim(omath, _arg)


# =========================================================
# Document building helpers
# =========================================================

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


CITATION_RE = re.compile(r'(\{[^}]+\}|\[\d+(?:[-,\u2013]\d+)*\])')


def _add_run_for_part(paragraph, part, base_size=Pt(12), bold=False, italic=False):
    if part.startswith('{') and part.endswith('}'):
        run = paragraph.add_run(part[1:-1])
        run.font.superscript = True
        run.font.size = Pt(9)
        return run
    if part.startswith('[') and part.endswith(']'):
        # Citation: keep brackets in normal size, superscript the numbers inside.
        paragraph.add_run('[').font.size = base_size
        inner = part[1:-1]
        run_inner = paragraph.add_run(inner)
        run_inner.font.superscript = True
        run_inner.font.size = Pt(9)
        paragraph.add_run(']').font.size = base_size
        return run_inner
    run = paragraph.add_run(part)
    run.font.size = base_size
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run


def add_paragraph(doc, text, style='Normal', bold=False, italic=False,
                  alignment=None, space_after=None, space_before=None):
    # Metrics placeholders such as {fano_auc} are replaced from the results files.
    text = text.format(**_METRICS)
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    parts = CITATION_RE.split(text)
    for part in parts:
        if part:
            _add_run_for_part(p, part, base_size=Pt(12), bold=bold, italic=italic)
    return p


def add_reference(doc, text, space_after=None):
    p = doc.add_paragraph(style='Normal')
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(12)
    return p


def add_figure(doc, img_path, caption, width=Inches(5.5)):
    caption = caption.format(**_METRICS)
    if not img_path.exists():
        add_paragraph(doc, "[MISSING FIGURE: " + img_path.name + "]")
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(12)
    run = p_img.add_run()
    run.add_picture(str(img_path), width=width)
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(12)
    parts = CITATION_RE.split(caption)
    for part in parts:
        if part:
            _add_run_for_part(p_cap, part, base_size=Pt(9))
    return p_cap


def add_table(doc, headers, data, caption=None):
    if caption:
        caption = caption.format(**_METRICS)
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(12)
        p_cap.paragraph_format.space_after = Pt(6)
        parts = CITATION_RE.split(caption)
        for part in parts:
            if part:
                run_cap = _add_run_for_part(p_cap, part, base_size=Pt(9))
                run_cap.italic = True
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = val
            for p in table.rows[row_idx + 1].cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


# =========================================================
def load_results():
    """Load evaluation and demo metrics from results/*.json."""
    results_dir = SCRIPT_DIR.parent / 'results'
    with open(results_dir / 'evaluation_summary.json') as f:
        summary = json.load(f)
    with open(results_dir / 'demo_summary.json') as f:
        demo = json.load(f)

    def fmt(v):
        if isinstance(v, str):
            return v
        # Keep mean values consistent with the published strings (3 decimals).
        return f"{v:.3f}"

    def pct(v):
        return f"{100.0 * float(v):.1f}"

    m = summary['methods']
    return {
        'n_recordings': summary['n_recordings'],
        'n_valid': m['fano_filter']['n_valid'],
        # Table rows
        'temporal_nrr': m['temporal_filter']['nrr'],
        'temporal_spr': m['temporal_filter']['spr'],
        'temporal_f1': m['temporal_filter']['f1'],
        'temporal_auc': m['temporal_filter']['auc'],
        'pi_nrr': m['pi_dc_dvs']['nrr'],
        'pi_spr': m['pi_dc_dvs']['spr'],
        'pi_f1': m['pi_dc_dvs']['f1'],
        'pi_auc': m['pi_dc_dvs']['auc'],
        'fano_nrr': m['fano_filter']['nrr'],
        'fano_spr': m['fano_filter']['spr'],
        'fano_f1': m['fano_filter']['f1'],
        'fano_auc': m['fano_filter']['auc'],
        # Scalar values used in the text
        'temporal_auc_value': fmt(m['temporal_filter']['auc_mean']),
        'pi_auc_value': fmt(m['pi_dc_dvs']['auc_mean']),
        'fano_auc_value': fmt(m['fano_filter']['auc_mean']),
        'fano_spr_value': fmt(m['fano_filter']['spr_mean']),
        'fano_nrr_value': fmt(m['fano_filter']['nrr_mean']),
        'temporal_nrr_pct': pct(m['temporal_filter']['nrr_mean']),
        'temporal_spr_pct': pct(m['temporal_filter']['spr_mean']),
        'fano_spr_pct': pct(m['fano_filter']['spr_mean']),
        'fano_nrr_pct': pct(m['fano_filter']['nrr_mean']),
        'pi_spr_pct': pct(m['pi_dc_dvs']['spr_mean']),
        'a5_mean_snr': f"{summary['a5_simulation']['mean_improvement']:.1f}",
        'a5_max_snr': f"{summary['a5_simulation']['max_improvement']:.1f}",
        # Demo
        'demo_events_in': demo['events_in'],
        'demo_events_residual': demo['events_residual'],
        'demo_removal_pct': demo['removal_pct'],
        'demo_threshold': demo['threshold'],
        'demo_signal_pixels': demo['signal_pixels'],
        'demo_total_pixels': demo['total_pixels'],
    }


# =========================================================
# Build the integrated manuscript
# =========================================================

def build_manuscript():
    global _METRICS
    _METRICS = load_results()

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # =========================================================
    # Title (blinded: no author info)
    # =========================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Noise removal for dynamic vision sensors:\n'
        'a physics-informed stochastic resonance framework'
    )
    run.font.size = Pt(16)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(
        '[Author details are provided on the separate title page]'
    )
    run.font.size = Pt(11)
    run.italic = True

    # =========================================================
    # Highlights
    # =========================================================
    add_heading(doc, 'Highlights', level=1)
    highlights = [
        '\u2022 Closed-form optimal noise-removal level \u03c1* = \u221a(1 \u2212 \u03b8\u00b2/\u03c3\u00b2) '
        'derived for threshold-based event detectors via stochastic resonance.',

        '\u2022 Five-parameter circuit-level DVS noise model (A5) enables a '
        'physics-informed noise inverse problem for event-based sensors.',

        '\u2022 Fano-factor-based filter achieves ROC-AUC = {fano_auc_value} and {fano_spr_pct}% signal '
        'preservation on {n_recordings} EBSSA recordings, substantially outperforming '
        'implemented temporal and neural baselines.',

        '\u2022 Six-tier calibration hierarchy including Cal-6 (satellite-trail '
        'verification) provides quantitative in-operation model checking.',

        '\u2022 Open pipeline demonstrates {demo_removal_pct}% noise removal with '
        'satellite trajectory recovery, supported by A5 simulation.',
    ]
    for h in highlights:
        add_paragraph(doc, h, space_after=4)

    # =========================================================
    # Abstract
    # =========================================================
    add_heading(doc, 'Abstract', level=1)
    add_paragraph(doc, (
        'Dynamic Vision Sensors (DVS) offer microsecond temporal resolution and '
        '>120 dB dynamic range for space situational awareness, but background '
        'activity noise dominates under low-light conditions, overwhelming faint '
        'astronomical signals. We address this engineering challenge through two '
        'complementary contributions. First, we derive a closed-form optimal noise '
        'model accuracy \u03c1* = \u221a(1 \u2212 \u03b8\u00b2/\u03c3\u00b2) for threshold-based event detectors '
        'by unifying stochastic resonance (SR) theory with covariate adjustment '
        '(ANCOVA). The result prescribes reducing effective noise to the SR optimum '
        'and no further\u2014establishing that indiscriminate noise removal is '
        'counterproductive. Second, we instantiate this framework for DVS '
        'astronomical observation via Physics-Informed DeepClean for DVS (PI-DC-DVS), '
        'integrating a five-parameter circuit-level noise model (A5) with Bayesian '
        'inference. Systematic evaluation on {n_recordings} recordings from the Event-Based '
        'Space Situational Awareness (EBSSA) dataset demonstrates that a '
        'Fano-factor-based noise filter achieves ROC-AUC = {fano_auc_value} with {fano_spr_pct}% '
        'signal preservation, substantially outperforming the implemented temporal-filter '
        'baseline (AUC = {temporal_auc_value}). A six-tier calibration hierarchy, including a '
        'satellite-trail verification tier, provides in-operation model checking. '
        'A proof-of-concept demonstration achieves {demo_removal_pct}% noise removal with clear '
        'satellite trajectory recovery, and A5-based simulation predicts {a5_mean_snr}\u00d7 mean '
        'SNR improvement\u2014quantitatively matching the SR framework\u2019s excess-noise-regime '
        'prediction. Monte Carlo simulations confirm all analytical results.'
    ))

    add_paragraph(doc, (
        'Keywords: dynamic vision sensor; event camera; noise inverse problem; '
        'stochastic resonance; measurement uncertainty; calibration; '
        'space situational awareness'
    ), italic=True, space_after=12)

    # =========================================================
    # 1. Introduction
    # =========================================================
    add_heading(doc, '1. Introduction', level=1)
    add_paragraph(doc, (
        'In threshold-based detectors, noise plays a dual role. At low levels it is '
        'innocuous; at high levels it overwhelms the signal. Between these extremes, '
        'stochastic resonance (SR) produces a counter-intuitive optimum where a finite '
        'noise level maximizes signal detection [1\u20133]. Separately, covariate '
        'adjustment\u2014as in analysis of covariance (ANCOVA) [4]\u2014reduces noise variance '
        'by modeling its dependence on observable auxiliary variables. These two '
        'traditions have evolved in isolation: SR theory identifies the optimal noise '
        'level but offers no prescription for reaching it; denoising methods remove '
        'noise without asking whether total removal is desirable.'
    ))
    add_paragraph(doc, (
        'Dynamic Vision Sensors (DVS) are neuromorphic sensors in which each pixel '
        'independently and asynchronously emits an event when the logarithmic '
        'intensity change exceeds a threshold [5,6]. DVS offer microsecond temporal '
        'resolution, >120 dB dynamic range, and sparse output\u2014properties that make '
        'them attractive for space situational awareness (SSA) [7\u20139] and fast optical '
        'astronomy [10]. Under low-light conditions, however, shot-noise-induced '
        'background activity (BA) dominates, overwhelming faint astronomical signals '
        '[11,12]. Conventional noise reduction treats this as a signal inverse '
        'problem; here we explore the complementary noise inverse problem, where the '
        'noise generation mechanism is modelled as a physical forward process and '
        'solved inversely to reconstruct and subtract noise.'
    ))
    add_paragraph(doc, (
        'This noise inverse paradigm has been successful in gravitational-wave (GW) '
        'astronomy, where DeepClean [13] and iDQ [14] use auxiliary witness channels '
        'to model and subtract non-stationary instrumental noise. Yet no equivalent '
        'pipeline exists for DVS astronomical observation, and the question of how '
        'much noise should be removed\u2014as opposed to the simpler question of how much '
        'can be removed\u2014has not been addressed.'
    ))
    add_paragraph(doc, (
        'We connect these perspectives by analyzing covariate adjustment in the '
        'presence of SR. For a threshold detector whose noise depends on observable '
        'covariates, we derive a closed-form optimal model accuracy as a function of '
        'the noise-to-threshold ratio. The answer is not \u201cremove everything you can '
        'model,\u201d but rather \u201creduce noise to the SR optimum and stop.\u201d We then '
        'instantiate this framework for DVS astronomical observation via a five-parameter '
        'circuit-level noise model, a Fano-factor-based filter, and a six-tier '
        'calibration hierarchy. Systematic evaluation on {n_recordings} EBSSA recordings '
        'demonstrates practical noise removal consistent with theoretical predictions.'
    ))
    add_paragraph(doc, (
        'The remainder of this paper is organized as follows. Section 2 develops '
        'the theoretical framework unifying SR with covariate adjustment. Section 3 '
        'describes DVS noise physics and the PI-DC-DVS noise-removal pipeline. Section 4 '
        'presents Monte Carlo simulations validating the theory. Section 5 reports '
        'systematic experimental evaluation on EBSSA data. Section 6 presents a '
        'proof-of-concept demonstration. Section 7 discusses engineering implications '
        'for event-based sensor systems and sensor design.'
    ))

    # =========================================================
    # 2. Theoretical framework
    # =========================================================
    add_heading(doc, '2. Theoretical framework', level=1)

    add_heading(doc, '2.1. Stochastic resonance in threshold detectors', level=2)
    add_paragraph(doc, (
        'Consider a threshold detector receiving x(t) = s(t) + n(t), where s(t) is a '
        'deterministic signal and n(t) is zero-mean Gaussian noise with variance \u03c3\u00b2. '
        'The output event stream is E(t) = 1 if |x(t)| > \u03b8, 0 otherwise, with '
        'threshold \u03b8 > 0. This encompasses DVS pixels [5,15], single-photon '
        'avalanche diodes [16], neuromorphic circuits [17], and Schmitt triggers. '
        'For a subthreshold sinusoidal signal s(t) = A sin(2\u03c0f\u2080t) with A < \u03b8, '
        'the two-state theory [1] gives the output SNR as:'
    ))
    add_display_equation(doc, eq_sr_snr, eq_num='1')
    add_paragraph(doc, (
        'Eq. (1) peaks at \u03c3* = \u03b8, independent of signal amplitude A (Fig. 1). '
        'Below \u03b8, threshold crossings are too rare; far above \u03b8, events are frequent '
        'but carry little signal modulation. At \u03c3 \u2248 \u03b8 the two effects balance, '
        'producing a maximally informative event stream.'
    ))

    # Fig 1: Conceptual schematic
    add_figure(doc, FIG_DIR_SR / 'fig1_schematic.png',
               'Fig. 1. Conceptual schematic of the covariate-adjusted stochastic '
               'resonance framework. (a) Threshold-based event detector: a weak '
               'periodic signal s(t) embedded in Gaussian noise triggers events when '
               '|x(t)| > \u03b8. (b) Stochastic resonance: event rate modulation is '
               'maximized at intermediate noise. (c) Covariate adjustment with noise '
               'model correlation \u03c1 narrows the residual distribution, shifting the '
               'operating point on the SR curve.',
               width=Inches(6.0))

    add_heading(doc, '2.2. Covariate adjustment as noise reduction', level=2)
    add_paragraph(doc, (
        'Suppose n(t) depends on observable covariates z(t) = (T, I_bg, \u03b8_mismatch, '
        '\u2026) through a model \u0148(t) = f(z(t); \u03b2). Subtracting this estimate gives '
        'x_adj(t) = s(t) + \u03b5(t), where \u03b5(t) = n(t) \u2212 \u0148(t) is the residual. '
        'If the model achieves correlation \u03c1 = Corr(\u0148, n), then '
        'Var(\u03b5) = (1 \u2212 \u03c1\u00b2)\u03c3\u00b2, so:'
    ))
    add_display_equation(doc, eq_sigma_eff, eq_num='2')
    add_paragraph(doc, (
        'The adjustment slides the operating point leftward along the SR curve by '
        'the factor \u221a(1 \u2212 \u03c1\u00b2). This is the measurement-theoretic connection: the '
        'noise model accuracy directly determines the effective noise level for '
        'threshold-based detection.'
    ))

    add_heading(doc, '2.3. Optimal noise model accuracy', level=2)
    add_paragraph(doc, (
        'Substituting \u03c3_eff into Eq. (1) and maximizing over \u03c1 at fixed \u03c3 yields '
        'the optimal model accuracy:'
    ))
    add_display_equation(doc, eq_rho_star, eq_num='3')
    add_paragraph(doc, (
        'with \u03c1* = 0 for \u03c3 \u2264 \u03b8. The two regimes have distinct physics (Fig. 2a). '
        'For \u03c3 \u2264 \u03b8 the detector sits at or below the SR peak; removing noise moves '
        'it further from the optimum, so \u03c1* = 0\u2014the noise is beneficial and should '
        'not be removed. For \u03c3 > \u03b8 the system is above the peak and adjustment '
        'should bring the effective noise exactly to the SR optimum: '
        '\u03c3_eff = \u03c3\u221a(1 \u2212 \u03c1*\u00b2) = \u03b8.'
    ))
    add_paragraph(doc, (
        'The resulting SNR gain over the unadjusted case is:'
    ))
    add_display_equation(doc, eq_snr_gain, eq_num='4')
    add_paragraph(doc, (
        'which scales as ~exp(2\u03c3\u00b2/\u03b8\u00b2) for \u03c3 \u226b \u03b8 (Fig. 2b). The exponential '
        'growth reflects the steep penalty of operating far above the SR peak, '
        'and the correspondingly large payoff of accurate noise modeling in '
        'high-noise environments\u2014precisely the regime of DVS astronomical observation.'
    ))

    # Fig 2: Optimal rho*
    add_figure(doc, FIG_DIR_SR / 'fig3_optimal_rho.png',
               'Fig. 2. (a) Optimal noise model accuracy \u03c1* versus input noise. '
               'Shaded yellow: SR regime (\u03c3 < \u03b8) where \u03c1* = 0 (noise is beneficial). '
               'Shaded blue: excess noise regime (\u03c3 > \u03b8) where '
               '\u03c1* = \u221a(1 \u2212 \u03b8\u00b2/\u03c3\u00b2). Dashed line: analytical prediction. '
               '(b) Peak SNR improvement at optimal \u03c1* grows exponentially, '
               'reaching ~100\u00d7 at \u03c3/\u03b8 = 4.',
               width=Inches(5.5))

    add_heading(doc, '2.4. Noise model accuracy and SNR improvement', level=2)
    add_paragraph(doc, (
        'Let \u03b1 denote the noise model accuracy (the fraction of noise variance '
        'explained by the model):'
    ))
    add_display_equation(doc, eq_accuracy, eq_num='5')
    add_paragraph(doc, (
        'The residual noise after subtraction is:'
    ))
    add_display_equation(doc, eq_residual_noise, eq_num='6')
    add_paragraph(doc, (
        'yielding an SNR improvement ratio:'
    ))
    add_display_equation(doc, eq_snr_improvement, eq_num='7')
    add_paragraph(doc, (
        'At \u03b1 = 0.9, this gives 10\u00d7 improvement; at \u03b1 = 0.99, 100\u00d7. The mapping '
        'between \u03b1 and \u03c1 is \u03c1 \u2248 \u03b1 for the DVS case, so the optimal accuracy '
        'from Eq. (3) directly constrains how much noise should be removed. '
        'Crucially, Eq. (4) shows that the gain through the SR nonlinearity is '
        'exponential in \u03c3\u00b2/\u03b8\u00b2, exceeding the linear 1/(1\u2212\u03b1) estimate used in '
        'gravitational-wave analyses.'
    ))

    # =========================================================
    # 3. DVS noise physics and methods
    # =========================================================
    add_heading(doc, '3. DVS noise physics and methods', level=1)

    add_heading(doc, '3.1. A5 parametric noise model', level=2)
    add_paragraph(doc, (
        'The circuit-level physics of DVS noise has been systematically '
        'characterised. Gra\u00e7a and Delbruck [11] proved that photon shot noise sets '
        'a fundamental lower bound at twice the shot noise level. McReynolds et al. '
        '[12] demonstrated alternating ON\u2194OFF polarity patterns in shot-noise events. '
        'Most importantly, Gra\u00e7a and Delbruck [18] introduced a physically realistic '
        'DVS pixel model incorporating first-passage-time stochastic event '
        'generation, achieving >1000\u00d7 computational speedup. A five-parameter '
        'analytical noise-rate expression (the A5 model) takes the form:'
    ))
    add_display_equation(doc, eq_a5_model, eq_num='8')
    add_paragraph(doc, (
        'where I_dark,ref is the reference dark current rate, '
        '\u03b1 \u2248 0.06\u20130.08 K\u207b\u00b9 is the temperature coefficient, \u0394T is the '
        'temperature offset, and \u03b2 is the background illuminance sensitivity. '
        'The five parameters (I_dark,ref, \u03b1, \u03b2, \u03b8_ON, \u03b8_OFF) are fitted per-pixel '
        'during offline calibration. The covariates z(t) = (T, I_bg, \u03b8_mismatch) '
        'predict per-pixel noise rates via Eq. (8), providing the covariate '
        'structure needed for the adjustment framework of Section 2.2.'
    ))

    add_heading(doc, '3.2. PI-DC-DVS algorithm', level=2)
    add_paragraph(doc, (
        'Physics-Informed DeepClean for DVS (PI-DC-DVS) operates in four phases '
        '(Fig. 3).'
    ))

    # Fig 3: Pipeline
    add_figure(doc, FIG_DIR_DVS / 'fig2_g3_pipeline_en.png',
               'Fig. 3. System architecture of the PI-DC-DVS noise inverse problem '
               'pipeline. Four stages: (1) noise forward model construction using the '
               'A5 pixel model and auxiliary channels; (2) Bayesian inverse problem '
               'solution; (3) residual event stream generation via probabilistic '
               'thinning; (4) calibration and verification including Cal-6 satellite '
               'trail calibration.',
               width=Inches(6.0))

    add_paragraph(doc, (
        'Phase 1 (Offline calibration): Record dark events (lens cap), flat-field '
        'events (integrating sphere), and thermal sweep events (\u0394T = \u00b15\u00b0C). '
        'Fit the A5 forward model to obtain per-pixel parameter maps via MAP '
        'estimation:'
    ))
    add_display_equation(doc, eq_map_estimation, eq_num='9')
    add_paragraph(doc, (
        'Phase 2 (Online inference): A Physics-Informed Neural Network predicts '
        'per-pixel noise rates. The network comprises: (a) physics model layer '
        '(fixed weights, A5 baseline), (b) auxiliary-channel coupling layer '
        '(MLP 64-32-1), and (c) spatio-temporal correlation layer '
        '(Conv2D 3\u00d73). Output:'
    ))
    add_display_equation(doc, eq_nn_output, eq_num='10')
    add_paragraph(doc, (
        'Per-event noise probability (following iDQ [14]):'
    ))
    add_display_equation(doc, eq_p_noise, eq_num='11')
    add_paragraph(doc, (
        'Phase 3 (Residual generation): Soft subtraction assigns weight '
        'w_i = 1 \u2212 P_noise(e_i); events with w_i > w_threshold are retained. '
        'Phase 4 (Adaptive updates): Monitor residual Poisson statistics; '
        'apply Kalman-filter-like drift correction.'
    ))

    add_heading(doc, '3.3. Fano factor as noise discriminant', level=2)
    add_paragraph(doc, (
        'The Fano factor serves as a physics-grounded test statistic exploiting '
        'the known Poisson character of DVS shot noise:'
    ))
    add_display_equation(doc, eq_fano, eq_num='12')
    add_paragraph(doc, (
        'where N_k is the event count in temporal bin k. Pure Poisson noise gives '
        'F \u2248 1; periodically modulated rates (astronomical signals) produce F \u226b 1. '
        'Pixels with F \u2264 F_thr (\u2248 2) are labelled noise-dominated and define the '
        'local noise rate \u03bb_noise(x, y). Per-event noise probability follows from '
        'Eq. (11); events with P_noise > 0.5 are classified as noise. This is '
        'the DVS realization of covariate adjustment: Eq. (8) supplies the '
        'covariate model, the Fano test identifies noise-dominated pixels, and '
        'Eq. (11) effects the adjustment.'
    ))

    add_heading(doc, '3.4. Six-tier calibration framework', level=2)
    add_paragraph(doc, (
        'DVS output event streams rather than frames, requiring purpose-designed '
        'calibration procedures. We propose six tiers (Table 1):'
    ))

    add_table(doc,
              ['Tier', 'Condition', 'Purpose', 'Pass criterion'],
              [
                  ['Cal-1', 'Dark (lens cap)', 'Pure noise reference', '\u03c7\u00b2/dof < 1.5'],
                  ['Cal-2', 'Thermal sweep', 'Temperature dependence', 'Residual < 10%'],
                  ['Cal-3', 'Flat-field', 'Shot noise statistics', '\u03b1_flat > 0.9'],
                  ['Cal-4', 'Dynamic patterns', 'Injection-recovery', 'AUC > 0.95'],
                  ['Cal-5', 'Simulated astro.', 'End-to-end pipeline', '\u0394m > 2 mag'],
                  ['Cal-6', 'Satellite trails', 'In-operation verification', 'Det. rate > 95%'],
              ],
              caption='Table 1. Six-tier calibration framework for DVS noise model validation.')

    add_paragraph(doc, (
        'Cal-6 repurposes satellite light trails\u2014conventionally regarded as light '
        'pollution\u2014as natural calibration sources. Artificial satellites have '
        'precisely predictable trajectories (TLE + SGP4 propagator), providing '
        'abundant, cost-free injection-recovery tests under real observing '
        'conditions. While CCD sensors saturate on bright trails, DVS records them '
        'quantitatively across >120 dB dynamic range [19]. This is especially '
        'timely given the proliferation of satellite constellations (Starlink, '
        'OneWeb) [20].'
    ))

    # =========================================================
    # 4. Numerical simulations
    # =========================================================
    add_heading(doc, '4. Numerical simulations', level=1)

    add_heading(doc, '4.1. Monte Carlo validation of SR theory', level=2)
    add_paragraph(doc, (
        'We test the analytical predictions with Monte Carlo (MC) simulations '
        '(Fig. 4). The signal is s(t) = A sin(2\u03c0f\u2080t) with A/\u03b8 = 0.3, '
        'f\u2080 = 5 Hz, dt = 1 ms, and N = 10^5 steps per trial.'
    ))

    # Fig 4: SR curves
    add_figure(doc, FIG_DIR_SR / 'fig2_sr_curves.png',
               'Fig. 4. Stochastic resonance curves for a threshold detector '
               '(A/\u03b8 = 0.3). Solid lines: analytical SNR from Eq. (1). Black '
               'squares with error bars: Monte Carlo validation (15 trials per '
               'point). Covariate adjustment (\u03c1 > 0) shifts the SR peak rightward.',
               width=Inches(4.5))

    add_paragraph(doc, (
        'Fig. 4 plots output SNR against input noise for several \u03c1 values. MC '
        'estimates agree with the analytical curve for \u03c1 = 0, peaking at '
        '\u03c3/\u03b8 \u2248 1. The adjusted curves shift rightward to '
        '\u03c3/\u03b8 \u2248 1/\u221a(1 \u2212 \u03c1\u00b2) as predicted.'
    ))

    add_heading(doc, '4.2. Detection probability and ROC analysis', level=2)
    add_paragraph(doc, (
        'Fig. 5 shows detection probability P_D and false alarm probability P_FA '
        'versus input noise (A/\u03b8 = 0.4). Both decrease with \u03c1 because adjustment '
        'suppresses the effective noise. Fig. 6 shows the ROC curves at '
        '\u03c3/\u03b8 = 1.5: increasing \u03c1 lifts the ROC curve well above the chance '
        'diagonal, demonstrating improved signal\u2013noise discrimination.'
    ))

    # Fig 5: Detection probability
    add_figure(doc, FIG_DIR_SR / 'fig4_detection_probability.png',
               'Fig. 5. (a) Detection probability P_D and (b) false alarm '
               'probability P_FA versus input noise for different noise model '
               'accuracies \u03c1 (A/\u03b8 = 0.4).',
               width=Inches(5.5))

    # Fig 6: ROC comparison
    add_figure(doc, FIG_DIR_SR / 'fig5_roc_comparison.png',
               'Fig. 6. ROC curves at \u03c3_n/\u03b8 = 1.5 (excess noise regime). '
               'Covariate adjustment (\u03c1 = 0.95) achieves near-ideal separation.',
               width=Inches(3.8))

    add_heading(doc, '4.3. A5-based noise rate simulation', level=2)
    add_paragraph(doc, (
        'Using the A5 model, we simulate noise rates and SNR improvements across '
        'the temperature\u2013illuminance parameter space '
        '(T \u2208 [10, 65]\u00b0C, I_bg \u2208 [0.1, 1000] lux). Fig. 7 shows the predicted '
        'noise rate map, SNR improvement factor, and temperature dependence. '
        'The simulation predicts a mean SNR improvement of {a5_mean_snr}\u00d7 (max {a5_max_snr}\u00d7) at '
        '90% noise model accuracy (\u03b1 = 0.9).'
    ))

    # Fig 7: A5 simulation
    add_figure(doc, FIG_DIR_DVS / 'fig6_a5_simulation.png',
               'Fig. 7. A5-based noise rate simulation. (a) Predicted noise event '
               'rate [evt/s/pix]; (b) SNR improvement factor at 90% model accuracy; '
               '(c) SNR vs. temperature at fixed illuminance comparing methods.',
               width=Inches(6.0))

    # =========================================================
    # 5. Experimental evaluation
    # =========================================================
    add_heading(doc, '5. Experimental evaluation', level=1)

    add_heading(doc, '5.1. Dataset and methods', level=2)
    add_paragraph(doc, (
        'We use the public EBSSA dataset [7]: 236 recordings from DAVIS240C sensors '
        'observing satellites and stars, with 572 labelled objects. We select '
        '{n_recordings} recordings spanning both sensor configurations (180\u00d7240 and '
        '240\u00d7304 pixels). Three methods that require only event coordinates and timestamps '
        'are compared:'
    ))
    add_paragraph(doc, (
        '(1) Fano filter (proposed): Physics-based noise-removal approach using '
        'the Fano factor as a Poisson discriminant. '
        '(2) PI-DC-DVS NN (proposed, simplified): Three-layer neural network, '
        'self-supervised on noise-dominated pixels. '
        '(3) Temporal filter (baseline) [21]: Spatio-temporal neighbourhood filter. '
        'Established learning-based event denoisers that require intensity frames '
        '(e.g. EDnCNN, motion-compensated filters, bilateral event filters) are not '
        'included here because the public EBSSA release does not provide those '
        'auxiliary data; a dedicated SOTA benchmark on a matched dataset is under way.'
    ))

    add_heading(doc, '5.2. Results', level=2)
    add_paragraph(doc, (
        'Table 2 summarises the systematic evaluation results across all '
        '{n_recordings} recordings (valid pairs: {n_valid}).'
    ))

    add_table(doc,
              ['Method', 'NRR', 'SPR', 'F1', 'AUC'],
              [
                  ['Temporal filter', _METRICS['temporal_nrr'], _METRICS['temporal_spr'],
                   _METRICS['temporal_f1'], _METRICS['temporal_auc']],
                  ['PI-DC-DVS NN', _METRICS['pi_nrr'], _METRICS['pi_spr'],
                   _METRICS['pi_f1'], _METRICS['pi_auc']],
                  ['Fano filter', _METRICS['fano_nrr'], _METRICS['fano_spr'],
                   _METRICS['fano_f1'], _METRICS['fano_auc']],
              ],
              caption='Table 2. Systematic evaluation results (mean \u00b1 std) across '
                      '{n_recordings} EBSSA recordings. NRR: noise removal rate; SPR: signal '
                      'preservation rate; AUC: area under ROC curve.')

    add_paragraph(doc, (
        'The Fano filter achieves the best balance between noise removal and signal '
        'preservation (Fig. 8), with AUC = {fano_auc_value} substantially exceeding both the '
        'temporal filter (AUC = {temporal_auc_value}) and the simplified PI-DC-DVS NN '
        '(AUC = {pi_auc_value}). The temporal filter achieves the highest raw noise removal '
        '({temporal_nrr_pct}%) but destroys most signal (SPR = {temporal_spr_pct}%), disqualifying it for '
        'faint-object detection.'
    ))

    # Fig 8: Evaluation
    add_figure(doc, FIG_DIR_DVS / 'fig5_systematic_evaluation.png',
               'Fig. 8. Systematic evaluation of three denoising methods on {n_recordings} EBSSA '
               'recordings. Four-panel boxplot: (a) Noise Removal Rate, '
               '(b) Signal Preservation Rate, (c) F1 Score, (d) ROC-AUC. '
               'The Fano filter achieves the best overall balance.',
               width=Inches(5.5))

    add_heading(doc, '5.3. Interpretation in the SR framework', level=2)
    add_paragraph(doc, (
        'DVS astronomical observations lie deep in the excess-noise regime '
        '(\u03c3 \u226b \u03b8): dark-current rates dominate signal rates by orders of magnitude. '
        'The A5 + Fano model achieves \u03c1 \u2248 0.7\u20130.9, which Fig. 2b predicts '
        'should yield {a5_mean_snr}\u2013{a5_max_snr}\u00d7 SNR improvement\u2014matching the measured {a5_mean_snr}\u00d7. '
        'Fig. 9 maps the experimental results onto the SR framework: the '
        'Fano filter (covariate adjustment) achieves ROC-AUC = {fano_auc_value}, far '
        'exceeding temporal filtering, while preserving {fano_spr_pct}% of signal. '
        'That noise is not eliminated entirely (NRR = {fano_nrr_value}) accords with '
        'the prediction that over-removal past the SR optimum is counterproductive.'
    ))

    # Fig 9: DVS application in SR framework
    add_figure(doc, FIG_DIR_SR / 'fig7_dvs_application.png',
               'Fig. 9. DVS results in the SR framework. (a) ROC-AUC comparison: '
               'the Fano filter (covariate adjustment) achieves {fano_auc_value}, far exceeding '
               'temporal filtering. (b) NRR vs SPR trade-off: the Fano filter '
               'preserves {fano_spr_pct}% of signal while removing {fano_nrr_pct}% of noise.',
               width=Inches(5.5))

    add_paragraph(doc, (
        'The neural network (AUC = {pi_auc_value}), despite greater flexibility, lacks '
        'access to physics-informed covariates and cannot match the Fano filter. '
        'This illustrates a key prediction: what matters is the fidelity of the '
        'noise model (\u03c1), not the complexity of the method.'
    ))

    # =========================================================
    # 6. Proof-of-concept demonstration
    # =========================================================
    add_heading(doc, '6. Proof-of-concept demonstration', level=1)
    add_paragraph(doc, (
        'From {demo_events_in} input events, probabilistic thinning with threshold '
        '\u03c4 = {demo_threshold} yields {demo_events_residual} residual events\u2014a {demo_removal_pct}% noise removal rate. '
        'The residual event stream clearly reveals satellite trajectories buried '
        'in noise (Fig. 10).'
    ))

    # Fig 10: Demo
    add_figure(doc, FIG_DIR_DVS / 'fig3_noise_inverse_demo.png',
               'Fig. 10. Proof-of-concept on EBSSA Recording #0. (a) Raw event '
               'accumulation ({demo_events_in} events); (b) estimated noise rate map; '
               '(c) per-event noise probability distribution showing bimodal '
               'separation; (d) residual events after {demo_removal_pct}% noise removal with '
               'satellite trajectory clearly visible.',
               width=Inches(5.5))

    add_paragraph(doc, (
        'The Fano factor spatial map (Fig. 11a) shows clear separation between '
        'noise-dominated pixels (F \u2248 1) and signal-containing pixels (F \u226b 1). '
        'Signal candidate pixels ({demo_signal_pixels} out of {demo_total_pixels}) concentrate along satellite '
        'tracks. The detection limit improvement scales as:'
    ))
    add_display_equation(doc, eq_detection_limit, eq_num='13')
    add_paragraph(doc, (
        'giving \u0394m > 2.5 mag at \u03b1 = 0.9.'
    ))

    # Fig 11: SNR analysis
    add_figure(doc, FIG_DIR_DVS / 'fig4_sn_improvement.png',
               'Fig. 11. SNR improvement analysis. (a) Fano factor spatial map: '
               'noise-dominated pixels (blue, F \u2248 1) vs. signal-containing pixels '
               '(red, F \u226b 1); (b) temporal dynamics of event rate vs. noise model; '
               '(c) per-pixel SNR distribution before and after subtraction.',
               width=Inches(6.0))

    # =========================================================
    # 7. Discussion
    # =========================================================
    add_heading(doc, '7. Discussion', level=1)

    add_heading(doc, '7.1. Engineering implications', level=2)
    add_paragraph(doc, (
        'From an engineering perspective, \u03c1* = \u221a(1 \u2212 \u03b8\u00b2/\u03c3\u00b2) is a '
        'noise-removal design criterion for threshold-based sensors: the target '
        'residual noise level is \u03c3_eff = \u03b8, not \u03c3_eff \u2192 0. The criterion is a '
        'function of the operating noise environment (\u03c3/\u03b8), not of the sensor\u2019s '
        'intrinsic resolution. Reducing noise below the SR optimum can degrade '
        'signal detection; therefore noise-removal systems should report and '
        'control \u03c1 explicitly, analogous to a tunable gain in a feedback loop.'
    ))

    add_heading(doc, '7.2. Sensor design implications', level=2)
    add_paragraph(doc, (
        'A practical corollary is that threshold sensors should be co-designed with '
        'noise models. Hardware noise reduction is costly; if an accurate covariate '
        'model is available, the sensor can tolerate higher raw noise and rely on '
        'post-hoc adjustment. The design target becomes \u03c3_eff = \u03b8 after adjustment, '
        'not \u03c3 = \u03b8 in hardware. This applies to DVS, single-photon detectors [16], '
        'and related threshold devices.'
    ))

    add_heading(doc, '7.3. Cal-6: satellite trails as verification sources', level=2)
    add_paragraph(doc, (
        'Cal-6 turns a common operational nuisance into an in-operation '
        'verification source: satellite constellations provide continuous, '
        'cost-free calibration signals with ephemeris-based trajectories. For DVS '
        'observation, satellite trails become practical verification assets '
        'rather than liabilities, enabling routine field checks of the noise model '
        'without dedicated calibration hardware.'
    ))

    add_heading(doc, '7.4. Connection to forbidden-interval theorems', level=2)
    add_paragraph(doc, (
        'The forbidden-interval theorem [22,23] gives necessary and sufficient '
        'conditions for SR in non-Gaussian noise. Covariate adjustment reshapes the '
        'effective distribution; our expression \u03c1* provides a constructive '
        'prescription complementing the theorem\u2019s existential characterization.'
    ))

    add_heading(doc, '7.5. Broader applicability', level=2)
    add_paragraph(doc, (
        'The principle extends to any measurement system combining a threshold '
        'nonlinearity with structured noise: neural spike detection [24], quantum '
        'key distribution [25], radar target detection [26], and ion-channel sensing '
        '[27]. In each case the prescription is: reduce noise to the SR optimum, '
        'and accept the residual as beneficial.'
    ))

    add_heading(doc, '7.6. Limitations and future work', level=2)
    add_paragraph(doc, (
        'Limitations and mitigation strategies. '
        '(i) Theoretical scope: the SR derivation assumes a Gaussian subthreshold '
        'noise and a sinusoidal signal; Sec. 2.2 and the forbidden-interval '
        'framework [22] show how this extends to non-Gaussian noise with finite '
        'mean. '
        '(ii) Poisson vs. Gaussian: DVS dark-current shot noise is Poisson; the '
        'A5 model [18] expresses the total pixel current as a Poisson-driven '
        'temporal derivative, and the Fano factor naturally classifies Poisson '
        'background activity (F \u2248 1) against bursty signal (F \u226b 1). '
        '(iii) SOTA comparison: the present comparison is restricted to methods '
        'that run on the public EBSSA event data alone; EDnCNN and other '
        'learning-based denoisers that require intensity frames or dedicated '
        'training sets will be evaluated in an expanded benchmark. '
        '(iv) Calibration: Cal-6 currently uses a small number of satellite '
        'transits; a systematic Starlink-based verification campaign is planned. '
        '(v) Hardware and dataset extension: validation on SciDVS [28] hardware and '
        'the DVSNOISE20 [29] benchmark is planned. '
        '(vi) Generalisability: the derivation holds for any threshold detector with '
        'structured noise; the DVS implementation is the primary engineering '
        'demonstration.'
    ))

    # =========================================================
    # 8. Conclusions
    # =========================================================
    add_heading(doc, '8. Conclusions', level=1)
    add_paragraph(doc, (
        'We have presented an engineering framework for noise removal in '
        'threshold-based event sensors, connecting stochastic resonance with the '
        'noise inverse problem for dynamic vision sensors. The main contributions are:'
    ))
    conclusions = [
        '(1) A closed-form optimal noise-removal level '
        '\u03c1* = \u221a(1 \u2212 \u03b8\u00b2/\u03c3\u00b2) for threshold-based event detectors, prescribing '
        'noise reduction to the SR optimum rather than total removal.',

        '(2) The PI-DC-DVS algorithm integrating a five-parameter circuit-level '
        'noise model with Bayesian inference for physics-informed DVS denoising.',

        '(3) A Fano-factor-based filter achieving AUC = {fano_auc_value} on {n_recordings} EBSSA recordings, '
        'substantially outperforming the implemented temporal-filter baseline '
        '(AUC = {temporal_auc_value}) and the simplified PI-DC-DVS NN '
        '(AUC = {pi_auc_value}), while preserving {fano_spr_pct}% of signal.',

        '(4) A six-tier calibration framework with Cal-6 repurposing satellite '
        'trails as natural calibration sources for in-operation verification.',

        '(5) A proof-of-concept achieving {demo_removal_pct}% noise removal with satellite '
        'trajectory recovery, and A5-based simulation predicting {a5_mean_snr}\u00d7 mean SNR '
        'improvement\u2014consistent with the theoretical excess-noise-regime prediction.',

        '(6) Monte Carlo simulations confirming all analytical results, including '
        'the exponential SNR gain ~exp(2\u03c3\u00b2/\u03b8\u00b2) in the excess-noise regime.',
    ]
    for c in conclusions:
        add_paragraph(doc, c, space_after=4)

    add_paragraph(doc, (
        'The principle\u2014model the noise, subtract to the SR optimum, accept the '
        'residual as beneficial\u2014provides an engineering design criterion for '
        'noise-removal systems in any threshold-based detector operating in '
        'structured noise.'
    ))

    # =========================================================
    # Data availability
    # =========================================================
    add_heading(doc, 'Data availability', level=1)
    add_paragraph(doc, (
        'The EBSSA dataset is publicly available via the Tonic library [7]. '
        'All code, figure-generation scripts, evaluation code, and submission '
        'artifacts are publicly mirrored from the wip repository according to the '
        'sync-to-repos workflow.'
    ))

    # =========================================================
    # Declaration of competing interest
    # =========================================================
    add_heading(doc, 'Declaration of competing interest', level=1)
    add_paragraph(doc, (
        'The authors declare that they have no known competing financial interests '
        'or personal relationships that could have appeared to influence the work '
        'reported in this paper.'
    ))

    # =========================================================
    # Acknowledgments
    # =========================================================
    add_heading(doc, 'Acknowledgments', level=1)
    add_paragraph(doc, '[To be added.]', italic=True)

    # =========================================================
    # CRediT author statement
    # =========================================================
    add_heading(doc, 'CRediT author statement', level=1)
    add_paragraph(doc, '[To be added: Conceptualization, Methodology, Software, '
                  'Validation, Formal analysis, Investigation, Data curation, '
                  'Writing \u2013 original draft, Writing \u2013 review & editing, '
                  'Visualization.]', italic=True)

    # =========================================================
    # References (Results in Engineering: numbered [N], Vancouver order)
    # =========================================================
    add_heading(doc, 'References', level=1)
    references = [
        "[1] L. Gammaitoni, P. Hänggi, P. Jung, F. Marchesoni, Stochastic resonance, Rev. Mod. Phys. 70 (1998) 223–269.",
        "[2] M.D. McDonnell, D. Abbott, What is stochastic resonance? Definitions, misconceptions, debates, and its relevance to biology, PLoS Comput. Biol. 5 (2009) e1000348.",
        "[3] A.R. Bulsara, L. Gammaitoni, Tuning in to noise, Phys. Today 49 (1996) 39–45.",
        "[4] G.W. Snedecor, W.G. Cochran, Statistical Methods, eighth ed., Iowa State University Press, Ames, 1989.",
        "[5] G. Gallego, T. Delbrück, G. Orchard, C. Bartolozzi, B. Taba, A. Censi, S. Leutenegger, A.J. Davison, J. Conradt, K. Daniilidis, D. Scaramuzza, Event-based vision: A survey, IEEE Trans. Pattern Anal. Mach. Intell. 44 (2022) 154–180.",
        "[6] T. Delbruck, R. Graça, M. Paluch, Feedback control of event cameras, in: Proc. IEEE/CVF CVPR Workshops, 2021, pp. 1324–1332.",
        "[7] S. Afshar, A.P. Nicholson, A. van Schaik, G. Cohen, Event-based object detection and tracking for space situational awareness, preprint arXiv:1911.08730, 2019.",
        "[8] N. Ralph, D. Joubert, A. Jolley, S. Afshar, N. Tothill, A. van Schaik, G. Cohen, Real-time event-based unsupervised feature consolidation and tracking for space situational awareness, Front. Neurosci. 16 (2022) 821157.",
        "[9] K. Kamiński, G. Cohen, T. Delbruck, M. Żołnowski, M. Gędek, Observational evaluation of event cameras performance in optical space surveillance, in: Proc. 1st NEO and Debris Detection Conference, ESA, 2019.",
        "[10] J. Hoang, Neuromorphic cameras for ground-based atmospheric Cherenkov telescopes, preprint arXiv:2310.16321, 2023.",
        "[11] R. Graça, T. Delbruck, Unraveling the paradox of intensity-dependent DVS pixel noise, preprint arXiv:2109.08640, 2021.",
        "[12] B. McReynolds, R. Graça, T. Delbruck, Exploiting alternating DVS shot noise event pair statistics to reduce background activity, preprint arXiv:2304.03494, 2023.",
        "[13] G. Vajente, Y. Huang, M. Isi, J.C. Driggers, J.S. Kissel, M.J. Szczepanczyk, S. Vitale, Machine-learning nonstationary noise out of gravitational-wave detectors, Phys. Rev. D 101 (2020) 042003.",
        "[14] R. Essick, P. Godwin, C. Hanna, L. Blackburn, E. Katsavounidis, iDQ: Statistical inference of non-Gaussian noise with auxiliary degrees of freedom in gravitational-wave detectors, Mach. Learn.: Sci. Technol. 2 (2021) 015004.",
        "[15] C. Posch, T. Serrano-Gotarredona, B. Linares-Barranco, T. Delbruck, Retinomorphic event-based vision sensors: Bioinspired cameras with spiking output, Proc. IEEE 102 (2014) 1470–1484.",
        "[16] R.H. Hadfield, Single-photon detectors for optical quantum information applications, Nat. Photonics 3 (2009) 696–705.",
        "[17] G. Indiveri, B. Linares-Barranco, T.J. Hamilton, et al., Neuromorphic silicon neuron circuits, Front. Neurosci. 5 (2011) 73.",
        "[18] R. Graça, T. Delbruck, Towards a physically realistic computationally efficient DVS pixel model, preprint arXiv:2505.07386, 2025.",
        "[19] R. Cao, D. Galor, A. Kohli, J.L. Yates, L. Waller, Noise2Image: noise-enabled static scene recovery for event cameras, Optica 12 (2025) 46–55.",
        "[20] J. McDowell, The low Earth orbit satellite population and impacts of the SpaceX Starlink constellation, Astrophys. J. Lett. 892 (2020) L36.",
        "[21] T. Delbruck, Frame-free dynamic digital vision, in: Proc. Intl. Symp. on Secure-Life Electronics, 2008, pp. 21–26.",
        "[22] B. Kosko, S. Mitaim, Stochastic resonance in noisy threshold neurons, Neural Netw. 16 (2003) 755–761.",
        "[23] S. Mitaim, B. Kosko, Adaptive stochastic resonance in noisy neurons based on mutual information, IEEE Trans. Neural Netw. 15 (2004) 1526–1540.",
        "[24] P. Hänggi, Stochastic resonance in biology, ChemPhysChem 3 (2002) 285–290.",
        "[25] N. Gisin, G. Ribordy, W. Tittel, H. Zbinden, Quantum cryptography, Rev. Mod. Phys. 74 (2002) 145–195.",
        "[26] M.A. Richards, J.A. Scheer, W.A. Holm, Principles of Modern Radar: Basic Principles, SciTech Publishing, 2010.",
        "[27] S.M. Bezrukov, I. Vodyanoy, Noise-induced enhancement of signal transduction across voltage-dependent ion channels, Nature 378 (1995) 362–364.",
        "[28] R. Graça, S. Zhou, B. McReynolds, T. Delbruck, SciDVS: A scientific dynamic vision sensor, ESSERC 2024, DOI:10.1109/esserc62670.2024.10719521.",
        "[29] R.W. Baldwin, M. Almatrafi, V. Asari, K. Hirakawa, Event probability mask (EPM) and event denoising convolutional neural network (EDnCNN), Proc. CVPR, 2020.",
    ]
    for ref in references:
        add_reference(doc, ref, space_after=3)

    # Save
    out_path = OUT_DIR / 'manuscript_rie.docx'
    doc.save(str(out_path))
    print(f"Manuscript saved: {out_path}")
    return out_path


if __name__ == '__main__':
    build_manuscript()
