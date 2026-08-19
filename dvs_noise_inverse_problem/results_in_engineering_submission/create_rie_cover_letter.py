#!/usr/bin/env python3
"""
Generate cover letter for Results in Engineering (Elsevier) submission.

This is a resubmission following rejection from Measurement.
Reviewer #1's comments pertained to a different manuscript (PSO-FCNN antenna-array
optimisation); the remaining reviewer comments have been addressed in the revised
text as noted below.
"""

import json
from docx import Document
from docx.shared import Inches, Pt
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent
RESULTS_DIR = OUT_DIR.parent / 'results'


def load_metrics():
    with open(RESULTS_DIR / 'evaluation_summary.json') as f:
        summary = json.load(f)
    with open(RESULTS_DIR / 'demo_summary.json') as f:
        demo = json.load(f)
    m = summary['methods']
    return {
        'n_recordings': summary['n_recordings'],
        'fano_auc': m['fano_filter']['auc_mean'],
        'fano_spr': m['fano_filter']['spr_mean'],
        'demo_removal': demo['removal_pct'],
    }


def build_cover_letter():
    metrics = load_metrics()
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Date
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run('[Date]')
    run.font.size = Pt(12)

    # Author details
    p = doc.add_paragraph()
    run = p.add_run('[Corresponding author name and affiliation]')
    run.font.size = Pt(12)

    # Addressee
    p = doc.add_paragraph()
    run = p.add_run(
        'Editor-in-Chief\n'
        'Results in Engineering\n'
        'Elsevier'
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run('Dear Editor,')
    run.font.size = Pt(12)

    # Body
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(
        'We submit the enclosed manuscript entitled "Noise removal for dynamic '
        'vision sensors: a physics-informed stochastic resonance framework" for '
        'consideration as a Full Research Paper in Results in Engineering.'
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        'This manuscript addresses the engineering problem of noise removal for '
        'event-based sensors operating in low-light conditions. We present a '
        'physics-informed noise removal framework that combines a five-parameter '
        'DVS pixel noise model (A5) with stochastic resonance (SR) theory. The '
        'framework answers how much noise a model should remove: reduce structured '
        'noise to the SR optimum rather than eliminating it, because over-removal '
        'degrades detection performance.'
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        f"We instantiate the framework for DVS astronomical observation, an emerging "
        f"engineering application for space situational awareness. A Fano-factor-based "
        f"filter and a six-tier calibration hierarchy provide practical tools for "
        f"noise removal and model verification. Systematic evaluation on the public "
        f"EBSSA dataset ({metrics['n_recordings']} recordings) demonstrates ROC-AUC = "
        f"{metrics['fano_auc']:.3f} with {metrics['fano_spr']*100:.1f}% signal "
        f"preservation, substantially outperforming the implemented temporal-filter "
        f"baseline. A proof-of-concept achieves {metrics['demo_removal']}% noise "
        f"removal with clear satellite trajectory recovery."
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        'The work is grounded in engineering because it: '
        '(1) addresses a concrete sensor-system problem\u2014removing structured noise '
        'from event-based imagers; (2) develops a physics-informed noise model and '
        'a computationally efficient filter; (3) provides an open reproducible '
        'pipeline and calibration hierarchy; and (4) demonstrates quantitative '
        'improvements on a public event-camera dataset.'
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        'This manuscript was previously submitted to Measurement (Elsevier) and '
        'received a reject decision. We have revised it for Results in Engineering: '
        'the narrative has been reframed around noise removal for an engineering '
        'readership. We explicitly note that Reviewer #1\u2019s comments did not pertain '
        'to this manuscript (they concerned a PSO-FCNN antenna-array optimisation '
        'study). The remaining reviewer concerns have been addressed as follows: '
        '(i) the Gaussian/Poisson consistency is now discussed in Sec. 7.6; '
        '(ii) the limited SOTA comparison is transparently framed in Sec. 5.1 and '
        'declared as planned extension work; '
        '(iii) the theoretical assumptions and calibration reliability (Cal-6) are '
        'discussed in Sec. 7.6; '
        '(iv) the engineering scope and claims are narrowed to avoid '
        'over-generalisation. '
        'All code and artifacts are publicly mirrored to support reproducibility. '
        'The authors declare no competing interests.'
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('Sincerely,')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('[Author name]\n[Affiliation]\n[Email]')
    run.font.size = Pt(12)
    run.italic = True

    out_path = OUT_DIR / 'cover_letter_rie.docx'
    doc.save(str(out_path))
    print(f"Cover letter saved: {out_path}")
    return out_path


if __name__ == '__main__':
    build_cover_letter()
