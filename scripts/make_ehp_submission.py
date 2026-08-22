#!/usr/bin/env python3
"""
Convert the base manuscript into an Environmental Health Perspectives (EHP)
submission package.

Re-uses make_manuscript.py so every reported number is read from the generated
result CSVs.  The post-processing steps enforce EHP formatting:
  - structured abstract (Background, Objectives, Methods, Results, Discussion)
  - keywords
  - Vancouver superscript citations with compressed ranges
  - EHP-style reference list
  - removal of the Research-in-context panel
  - EHP declaration headings (Author Contributions, Competing Interests, AI Disclosure)
  - separate figure files + legends-only manuscript
  - cover letter including the EHP Statement of Contribution and
    Environmental-Health Significance
"""
import os
import re
import shutil
import sys
import tempfile

from docx import Document
from docx.shared import Pt

# make_manuscript.py lives in the same scripts/ directory
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
BW = os.environ.get("FIGURES_BW") == "1"
SUFFIX = "_bw" if BW else ""
DEFAULT_FIG = os.path.join(ROOT, "output", "figures_bw" if BW else "figures")
DEFAULT_MAN = os.path.join(ROOT, "output", "manuscript_bw" if BW else "manuscript")
os.environ.setdefault("FIGURES_DIR", DEFAULT_FIG)
os.environ.setdefault("MANUSCRIPT_DIR", DEFAULT_MAN)

sys.path.insert(0, HERE)
import make_manuscript as mm  # noqa: E402
MAN = mm.MAN

# ---------------------------------------------------------------------------
# EHP formatting constants
# ---------------------------------------------------------------------------
EHP_KEYWORDS = (
    "heat; traffic crash; mortality; climate change; "
    "road safety; distributed-lag model"
)

EHP_REF_TEXT = {
    "dlnm": (
        "Gasparrini A, Armstrong B, Kenward MG. 2010. Distributed lag non-linear models. "
        "Stat Med 29(21):2224-2234, PMID: 20812303, 10.1002/sim.3940."
    ),
    "lancet": (
        "Gasparrini A, Guo Y, Hashizume M, Lavigne E, Zanobetti A, Schwartz J, et al. 2015. "
        "Mortality risk attributable to high and low ambient temperature: a multicountry observational study. "
        "Lancet 386(9991):369-375, PMID: 26003380, 10.1016/S0140-6736(14)62114-0."
    ),
    "basu": (
        "Basu R. 2009. High ambient temperature and mortality: a review of epidemiologic studies from 2001 to 2008. "
        "Environ Health 8:40, PMID: 19758453, 10.1186/1476-069X-8-40."
    ),
    "fars": (
        "National Highway Traffic Safety Administration. n.d. Fatality Analysis Reporting System (FARS). "
        "https://www.nhtsa.gov/research-data/fatality-analysis-reporting-system-fars "
        "[accessed 22 August 2026]."
    ),
    "ghcn": (
        "Menne MJ, Durre I, Vose RS, Gleason BE, Houston TG. 2012. "
        "An overview of the Global Historical Climatology Network-Daily database. "
        "J Atmos Ocean Technol 29(7):897-910, 10.1175/JTECH-D-11-00103.1."
    ),
    "cdc": (
        "Centers for Disease Control and Prevention. n.d. CDC WONDER Underlying Cause of Death database. "
        "https://wonder.cdc.gov/ [accessed 22 August 2026]."
    ),
    "npa": (
        "National Police Agency of Japan. n.d. Traffic accident statistics open data. "
        "https://www.npa.go.jp/publications/statistics/koutsuu/opendata/ "
        "[accessed 22 August 2026]."
    ),
    "attrib": (
        "Gasparrini A, Leone M. 2014. Attributable risk from distributed lag models. "
        "BMC Med Res Methodol 14:55, PMID: 24758509, 10.1186/1471-2288-14-55."
    ),
    "eia": (
        "US Energy Information Administration. n.d. Weekly finished motor gasoline product supplied (PET.WGFUPUS2.W). "
        "https://www.eia.gov/ [accessed 22 August 2026]."
    ),
    "fhwa": (
        "Federal Highway Administration. n.d. Traffic Volume Trends. "
        "https://www.fhwa.dot.gov/policyinformation/travel_monitoring/tvt.cfm "
        "[accessed 22 August 2026]."
    ),
    "fhwa_vm2": (
        "Federal Highway Administration. n.d. Highway Statistics VM-2: Annual state vehicle-miles travelled. "
        "https://www.fhwa.dot.gov/policyinformation/statistics.cfm "
        "[accessed 22 August 2026]."
    ),
    "census_pep": (
        "US Census Bureau. n.d. Population Estimates Program. Annual state population estimates. "
        "https://www.census.gov/programs-surveys/popest.html "
        "[accessed 22 August 2026]."
    ),
    "ipcc": (
        "IPCC. 2021. Climate Change 2021: The Physical Science Basis. "
        "Contribution of Working Group I to the Sixth Assessment Report of the Intergovernmental Panel on Climate Change. "
        "Cambridge: Cambridge University Press."
    ),
    "daanen": (
        "Daanen HAM, van de Vliert E, Huang X. 2003. "
        "Driving performance in cold, warm, and thermoneutral environments. "
        "Appl Ergon 34(6):597-602, PMID: 14559420, 10.1016/S0003-6870(03)00055-3."
    ),
    "liang2022": (
        "Liang M, Min M, Guo X, Song Q, Wang H, Li N, et al. 2022. "
        "The relationship between ambient temperatures and road traffic injuries: a systematic review and meta-analysis. "
        "Environ Sci Pollut Res 29(33):50647-50660, PMID: 35235122, 10.1007/s11356-022-19437-y."
    ),
    "liang2021_aap": (
        "Liang M, Zhao D, Wu Y, Ye P, Wang Y, Yao Z, et al. 2021. "
        "Short-term effects of ambient temperature and road traffic accident injuries in Dalian, Northern China: "
        "A distributed lag non-linear analysis. Accid Anal Prev 153:106057, PMID: 33647596, 10.1016/j.aap.2021.106057."
    ),
    "care": (
        "European Commission, Directorate-General for Mobility and Transport. n.d. "
        "Community database on road accidents (CARE). "
        "https://road-safety.transport.ec.europa.eu/european-road-safety-observatory/methodology-and-research/care-database_en "
        "[accessed 22 August 2026]."
    ),
}

EHP_COVER = """22 August 2026

The Editor-in-Chief, Environmental Health Perspectives

Dear Editor-in-Chief,

We submit for your consideration our manuscript, "Ambient heat as an under-recognised risk factor for US traffic-crash mortality: a distributed-lag analysis with road-safety implications", as a Research Article.

Statement of Contribution and Environmental-Health Significance:
Ambient heat is a well-established environmental exposure that increases all-cause and cardiovascular mortality, but its contribution to fatal road traffic crashes has been largely overlooked. Using only public data, this ecological time-series study shows that days hotter than the local seasonal norm carry an acute, same-day excess of US traffic-crash deaths. The excess survives adjustment for aggregate driving activity and remains in the primary VIF-screened full-controls model (VIF threshold 10; retaining state population as a log offset, daily precipitation and estimated wet-bulb globe temperature, and dropping state vehicle-miles travelled and the humidex and heat-index anomalies); a stricter VIF < 5 screen and an unscreened all-controls model are reported as sensitivity analyses. The heat-stress metrics, tested individually, yield variable cumulative estimates consistent with collinearity. The excess is concentrated in heat-exposed open-air road users (motorcyclists, pedestrians and cyclists), and is comparable in magnitude to all officially recorded direct-heat deaths. Because these deaths continue to be coded as ordinary traffic crashes, the heat contribution remains invisible to heat-mortality surveillance and road-safety statistics, suggesting an uncounted environmental-health burden that would grow under continued warming and could be targeted by heat-aware road safety messaging and shared-mobility heat-adaptation measures.

This work advances the environmental health literature by (1) quantifying a hidden, climate-sensitive mortality burden in a cause-of-death category normally outside the scope of heat research, (2) showing that the open-air road-user gradient is consistent with direct heat exposure, and (3) providing reproducible national estimates and projections to inform climate adaptation and transportation safety policy. The analysis is intentionally transparent: all data and code are public at https://github.com/bougtoir/heat-accidents-mortality, and the pipeline is fully reproducible (make all for data and figures, then make ehp for the manuscript and submission package), with no hard-coded results.

The manuscript is original, is not under consideration elsewhere, and all authors approve submission. We declare no competing interests. We confirm that this manuscript has not been posted to a preprint server.

We look forward to your consideration.

Yours sincerely,


Tatsuki Onishi, on behalf of all authors

Data Science and AI Innovation Research Promotion Center

Shiga University of Medical Science

Seta Tsukinowa-cho, Otsu, Shiga 520-2192, Japan

E-mail: bougtoir@gmail.com
"""


# ---------------------------------------------------------------------------
# docx helpers (reused / adapted from make_aap_submission.py)
# ---------------------------------------------------------------------------
def _heading_text(p):
    return (p.text or "").strip()


def _set_heading_text(p, text):
    p.text = text


def _remove_paragraph(p):
    p._element.getparent().remove(p._element)


def _compress_citation_text(nums_str):
    """Turn '1,2,4,5,7' into '1-2,4-5,7' using an en dash for consecutive runs."""
    nums = [int(x) for x in nums_str.split(",") if x.strip()]
    if not nums:
        return nums_str
    groups = []
    for n in nums:
        if groups and n == groups[-1][-1] + 1:
            groups[-1].append(n)
        else:
            groups.append([n])
    parts = []
    for g in groups:
        if len(g) >= 2:
            parts.append(f"{g[0]}–{g[-1]}")
        else:
            parts.append(str(g[0]))
    return ",".join(parts)


def _compress_citations(doc):
    """Compress consecutive superscript citation numbers to en-dash ranges."""
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.superscript and r.text and re.fullmatch(r"\d+(,\d+)", r.text):
                r.text = _compress_citation_text(r.text)


# ---------------------------------------------------------------------------
# Abstract / front-matter post-processing
# ---------------------------------------------------------------------------
def _make_ehp_abstract(doc, abstract_heading, body_ps):
    """Rebuild Summary as an EHP structured abstract + keywords."""
    sections = {}
    funding_p = None
    for p in body_ps:
        txt = p.text or ""
        if txt.startswith("Background"):
            sections["Background"] = re.sub(r"^Background\s*", "", txt)
        elif txt.startswith("Methods"):
            sections["Methods"] = re.sub(r"^Methods\s*", "", txt)
        elif txt.startswith("Findings"):
            sections["Results"] = re.sub(r"^Findings\s*", "", txt)
        elif txt.startswith("Interpretation"):
            sections["Discussion"] = re.sub(r"^Interpretation\s*", "", txt)
        elif txt.startswith("Funding"):
            funding_p = p
            sections["Funding"] = re.sub(r"^Funding\s*", "", txt)

    # Split the aim sentence out of Background to form Objectives.
    background = sections.get("Background", "")
    m = re.search(r"^(.*\.\s+)(We\s+(?:estimated|aimed|sought|assessed|quantified)[^\.]*\.)", background)
    if m:
        background_only = m.group(1).strip()
        objectives = m.group(2).strip()
    else:
        background_only = background
        objectives = ""

    # Clean out the old abstract body paragraphs.
    for p in body_ps:
        _remove_paragraph(p)

    abstract_heading.text = "Abstract"

    order = [
        ("Background", background_only),
        ("Objectives", objectives),
        ("Methods", sections.get("Methods", "")),
        ("Results", sections.get("Results", "")),
        ("Discussion", sections.get("Discussion", "")),
    ]

    prev = abstract_heading
    for label, text in order:
        if not text:
            continue
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: ")
        r.bold = True
        p.add_run(text)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(6)
        prev._element.addnext(p._element)
        prev = p

    # Keywords replace the former Funding paragraph.
    kw = doc.add_paragraph()
    r = kw.add_run("Keywords: ")
    r.bold = True
    kw.add_run(EHP_KEYWORDS)
    kw.paragraph_format.line_spacing = 2.0
    kw.paragraph_format.space_after = Pt(6)
    prev._element.addnext(kw._element)


def _remove_research_in_context(doc, heading, body_ps):
    _remove_paragraph(heading)
    for p in body_ps:
        _remove_paragraph(p)


def _rename_headings(doc):
    """Adjust declaration headings to EHP conventions and add Acknowledgments."""
    renames = {
        "Contributors": "Author Contributions",
        "Declaration of competing interests": "Competing Interests",
        "Declaration of generative AI use": "AI Disclosure",
        "Data availability": "Data Availability",
    }
    author_contributions_heading = None
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            txt = _heading_text(p)
            if txt in renames:
                _set_heading_text(p, renames[txt])
                if txt == "Contributors":
                    author_contributions_heading = p

    if author_contributions_heading is not None:
        ack_h = doc.add_paragraph()
        ack_h.style = doc.styles["Heading 1"]
        ack_h.add_run("Acknowledgments")
        ack_b = doc.add_paragraph("None declared.")
        ack_b.paragraph_format.space_after = Pt(6)
        author_contributions_heading._element.addprevious(ack_h._element)
        ack_h._element.addnext(ack_b._element)


def _fix_data_availability(doc):
    """The base manuscript mentions 'make aap'; update to 'make ehp' for this target."""
    for p in doc.paragraphs:
        if "make aap for the manuscript and submission package" in (p.text or ""):
            p.text = p.text.replace(
                "make aap for the manuscript and submission package",
                "make ehp for the manuscript and submission package",
            )


def _postprocess(src_path, dst_path):
    doc = Document(src_path)

    # 1. Citation style.
    _compress_citations(doc)

    # 2. Abstract and Research-in-context.
    abstract_heading = None
    abstract_body = []
    research_heading = None
    research_body = []

    collect = None
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            txt = _heading_text(p)
            if txt == "Summary":
                abstract_heading = p
                collect = "abstract"
                continue
            if txt == "Research in context":
                research_heading = p
                collect = "research"
                continue
            collect = None
            continue
        if collect == "abstract":
            abstract_body.append(p)
        elif collect == "research":
            research_body.append(p)

    if abstract_heading and abstract_body:
        _make_ehp_abstract(doc, abstract_heading, abstract_body)

    if research_heading:
        _remove_research_in_context(doc, research_heading, research_body)

    # 3. Declaration headings.
    _rename_headings(doc)

    # 4. Data availability wording.
    _fix_data_availability(doc)

    doc.save(dst_path)
    print("wrote", dst_path)


# ---------------------------------------------------------------------------
# Cover letter and submission bundle
# ---------------------------------------------------------------------------
def _build_cover_letter(path):
    doc = Document()
    mm.setup(doc)
    for block in EHP_COVER.strip().split("\n\n"):
        text = " ".join(block.splitlines())
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
    doc.save(path)
    print("wrote", path)


def _build_figures_stage():
    """Copy the existing submission figure files into a temporary staging folder."""
    dest_dir = tempfile.mkdtemp(prefix="ehp_figures_", dir=MAN)
    src = os.path.join(MAN, "submission_figures")
    if not os.path.isdir(src):
        raise FileNotFoundError(f"{src} not found; run 'make manuscript' first.")
    shutil.copytree(src, os.path.join(dest_dir, "ehp_submission_figures"))
    return dest_dir


def _build_submission_zip(fig_stage, suffix=""):
    zip_base = os.path.join(MAN, f"ehp_submission_package{suffix}")
    stage = zip_base + "_stage"
    if os.path.exists(stage):
        shutil.rmtree(stage)
    os.makedirs(stage)
    files = [
        f"heat_crash_mortality_ehp{suffix}.docx",
        f"heat_crash_mortality_ehp_legends{suffix}.docx",
        f"ehp_cover_letter{suffix}.docx",
        f"tables{suffix}.docx",
        f"figures{suffix}.pptx",
        f"strobe_checklist{suffix}.docx",
    ]
    for name in files:
        src = os.path.join(MAN, name)
        # Base make_manuscript outputs do not carry the _bw suffix; rename on copy.
        if not os.path.exists(src) and suffix:
            base_name = name.replace(suffix, "", 1)
            src = os.path.join(MAN, base_name)
        if os.path.exists(src):
            shutil.copyfile(src, os.path.join(stage, name))
    if fig_stage and os.path.isdir(fig_stage):
        for entry in os.listdir(fig_stage):
            entry_src = os.path.join(fig_stage, entry)
            if entry == "ehp_submission_figures" and suffix:
                entry_dst = os.path.join(stage, f"ehp_submission_figures{suffix}")
            else:
                entry_dst = os.path.join(stage, entry)
            if os.path.isdir(entry_src):
                shutil.copytree(entry_src, entry_dst)
            else:
                shutil.copyfile(entry_src, entry_dst)
    if os.path.exists(zip_base + ".zip"):
        os.remove(zip_base + ".zip")
    shutil.make_archive(zip_base, "zip", stage)
    shutil.rmtree(stage)
    print("wrote", zip_base + ".zip")


def main():
    os.makedirs(MAN, exist_ok=True)

    # Use EHP reference formatting while re-building the base manuscript.
    mm.REF_TEXT = EHP_REF_TEXT

    # Ensure the editable tables and figures PPTX are current.
    mm.build_pptx()
    mm.build_tables_docx()

    # Build base inline and legends-only manuscripts in a temporary directory.
    tmpdir = tempfile.mkdtemp(prefix="ehp_build_", dir=MAN)
    original_man = mm.MAN
    mm.MAN = tmpdir
    try:
        mm.build_manuscript("tmp_inline.docx", embed=True)
        mm.build_manuscript("tmp_legends.docx", embed=False)
    finally:
        mm.MAN = original_man

    inline_src = os.path.join(tmpdir, "tmp_inline.docx")
    legends_src = os.path.join(tmpdir, "tmp_legends.docx")

    ehp_inline = os.path.join(MAN, f"heat_crash_mortality_ehp{SUFFIX}.docx")
    ehp_legends = os.path.join(MAN, f"heat_crash_mortality_ehp_legends{SUFFIX}.docx")

    _postprocess(inline_src, ehp_inline)
    _postprocess(legends_src, ehp_legends)

    shutil.rmtree(tmpdir)

    _build_cover_letter(os.path.join(MAN, f"ehp_cover_letter{SUFFIX}.docx"))

    fig_stage = _build_figures_stage()
    try:
        _build_submission_zip(fig_stage, SUFFIX)
    finally:
        shutil.rmtree(fig_stage)


if __name__ == "__main__":
    main()
